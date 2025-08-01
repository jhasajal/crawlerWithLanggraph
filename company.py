from firecrawl import FirecrawlApp, ScrapeOptions
from docx import Document
import json 

app = FirecrawlApp(api_key="fc-745b0d56f5d04dc3aa8b16e3399bd580")

# Crawl a website:
crawl_result = app.crawl_url(
    'https://www.munichre.com/',
    limit=10,
    scrape_options=ScrapeOptions(formats=['markdown', 'html']),
)

# Write crawled data to a DOCX file
doc = Document()
doc.add_heading("Website Crawl Report: https://www.munichre.com/", 0)

doc.add_paragraph(json.dumps(crawl_result.model_dump(), indent=2, default=str))

pages = getattr(crawl_result, 'pages', [])
if pages:
    for idx, page in enumerate(pages, 1):
        url = getattr(page, 'url', 'N/A')
        markdown = getattr(page, 'markdown', None)
        html = getattr(page, 'html', None)
        doc.add_heading(f"{idx}. {url}", level=1)
        if markdown:
            doc.add_heading("Markdown Content", level=2)
            doc.add_paragraph(markdown)
        if html:
            doc.add_heading("HTML Content", level=2)
            doc.add_paragraph(html)
else:
    doc.add_paragraph("No pages found in crawl result.")

output_filename = "Website_Crawl_Report.docx"
doc.save(output_filename)
print(f"✅ DOCX file created: {output_filename}")