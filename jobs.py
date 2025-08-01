from firecrawl import FirecrawlApp
from docx import Document

app = FirecrawlApp(api_key="fc-745b0d56f5d04dc3aa8b16e3399bd580")

# Scrape a website:
scrape_result = app.scrape_url(
    'https://munichre-jobs.com/de/MunichRe',
    formats=['links'],
    actions=[
        {"type": "click", "selector": "a.page-link[data-new-page]"},
        {"type": "wait", "milliseconds": 7000},
        {"type": "scrape"}
    ]
)


# Write scraped links to a DOCX file
doc = Document()
doc.add_heading("Scraped Links", 0)

links = getattr(scrape_result, 'links', [])
if links:
    for link in links:
        doc.add_paragraph(link)
else:
    doc.add_paragraph("No links found.")

doc.save("Scraped_Links.docx")
print("✅ DOCX file created: Scraped_Links.docx")